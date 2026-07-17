from __future__ import annotations

import html
import io
import json
import logging
import os
import re
import ssl
import subprocess
import tempfile
import urllib.request
import zipfile
from base64 import b64encode
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse
from urllib.request import Request

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".html", ".htm", ".json", ".xml"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
ARCHIVE_EXTENSIONS = {".zip", ".rar", ".7z"}
DEFAULT_TESSERACT_CMD = Path(r"D:\Tesseract-OCR\tesseract.exe")
IMPORTANT_ARCHIVE_NAME_PARTS = (
    "тз",
    "tz",
    "brief",
    "бриф",
    "задани",
    "тех",
    "опис",
    "макет",
    "design",
    "figma",
    "screen",
    "скрин",
    "форма",
)


@dataclass(frozen=True)
class AttachmentRef:
    label: str
    url: str


@dataclass(frozen=True)
class AttachmentReport:
    label: str
    url: str
    local_path: str
    status: str
    summary: str
    kind: str
    opened_archive: bool
    ocr_scanned: bool


@dataclass(frozen=True)
class AttachmentProcessingResult:
    context: str
    reports: tuple[AttachmentReport, ...]


@dataclass(frozen=True)
class ArchiveEntryInfo:
    name: str
    size: int
    kind: str


@dataclass(frozen=True)
class ArchiveSelection:
    names: tuple[str, ...]
    used_ai: bool
    reason: str = ""


def build_attachment_context(
    attachments: tuple[str, ...],
    cookie: str = "",
    max_files: int = 3,
    max_bytes: int = 2_000_000,
    use_browser: bool = False,
    cdp_url: str = "http://127.0.0.1:9222",
    browser_profile_dir: str = "",
    lead_context: str = "",
    deepseek_api_key: str = "",
    deepseek_model: str = "deepseek-chat",
    openrouter_api_key: str = "",
    openrouter_base_url: str = "https://openrouter.ai/api/v1",
    openrouter_vision_model: str = "",
) -> str:
    """Download small readable attachments and return text for AI context."""
    return build_attachment_report(
        attachments,
        cookie=cookie,
        max_files=max_files,
        max_bytes=max_bytes,
        use_browser=use_browser,
        cdp_url=cdp_url,
        browser_profile_dir=browser_profile_dir,
        lead_context=lead_context,
        deepseek_api_key=deepseek_api_key,
        deepseek_model=deepseek_model,
        openrouter_api_key=openrouter_api_key,
        openrouter_base_url=openrouter_base_url,
        openrouter_vision_model=openrouter_vision_model,
    ).context


def build_attachment_report(
    attachments: tuple[str, ...],
    cookie: str = "",
    max_files: int = 3,
    max_bytes: int = 2_000_000,
    use_browser: bool = False,
    cdp_url: str = "http://127.0.0.1:9222",
    browser_profile_dir: str = "",
    output_dir: str | Path | None = None,
    lead_context: str = "",
    deepseek_api_key: str = "",
    deepseek_model: str = "deepseek-chat",
    openrouter_api_key: str = "",
    openrouter_base_url: str = "https://openrouter.ai/api/v1",
    openrouter_vision_model: str = "",
) -> AttachmentProcessingResult:
    """Download attachments, save readable originals, and return AI + UI metadata."""
    if not attachments:
        return AttachmentProcessingResult(context="", reports=())

    blocks: list[str] = ["ФАЙЛЫ/ТЗ:"]
    reports: list[AttachmentReport] = []
    for raw in attachments[:max_files]:
        ref = parse_attachment(raw)
        local_path = ""
        try:
            content = _download_with_fallback(
                ref,
                cookie=cookie,
                max_bytes=max_bytes,
                use_browser=use_browser,
                cdp_url=cdp_url,
                browser_profile_dir=browser_profile_dir,
            )
            local_path = _save_attachment_file(ref, content, output_dir) if output_dir is not None else ""
            status, extracted = inspect_attachment(
                ref,
                content,
                max_bytes=max_bytes,
                lead_context=lead_context,
                deepseek_api_key=deepseek_api_key,
                deepseek_model=deepseek_model,
                openrouter_api_key=openrouter_api_key,
                openrouter_base_url=openrouter_base_url,
                openrouter_vision_model=openrouter_vision_model,
            )
        except Exception as exc:
            logger.warning("Failed to read attachment %s: %s", ref.url, exc)
            status = "не скачан"
            extracted = f"Не удалось скачать или прочитать файл: {exc}"
        report = AttachmentReport(
            label=ref.label,
            url=ref.url,
            local_path=local_path,
            status=status,
            summary=_shorten(extracted, 2500),
            kind=_attachment_kind(ref),
            opened_archive="архив открыт" in status,
            ocr_scanned="OCR" in status or "vision" in status.lower(),
        )
        reports.append(report)
        blocks.append(
            "\n".join(
                [
                    f"- {ref.label}",
                    f"  Ссылка: {ref.url}",
                    f"  Статус: {status}",
                    f"  Кратко: {report.summary}",
                ]
            )
        )
    return AttachmentProcessingResult(context="\n\n".join(blocks), reports=tuple(reports))


def parse_attachment(raw: str) -> AttachmentRef:
    label, sep, url = raw.partition(": ")
    if not sep:
        url = raw.strip()
        label = url.rsplit("/", 1)[-1] or "attachment"
    return AttachmentRef(label=_clean_label(label), url=url.strip())


def _clean_label(value: str) -> str:
    value = re.sub(r"<[^>]+>", " ", html.unescape(value))
    return " ".join(value.split()).strip()


def _save_attachment_file(ref: AttachmentRef, content: bytes, output_dir: str | Path) -> str:
    directory = Path(output_dir)
    directory.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(ref.label)
    if not _extension(filename):
        ext = _extension(ref.url)
        filename = f"{filename}{ext}" if ext else filename
    path = _unique_path(directory / filename)
    path.write_bytes(content)
    return str(path)


def _safe_filename(value: str) -> str:
    name = Path(urlparse(value).path).name if re.match(r"https?://", value, re.IGNORECASE) else value
    name = _clean_label(name) or "attachment"
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name).strip(" ._")
    return name[:120] or "attachment"


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem or "attachment"
    suffix = path.suffix
    for index in range(2, 1000):
        candidate = path.with_name(f"{stem}-{index}{suffix}")
        if not candidate.exists():
            return candidate
    raise FileExistsError(f"слишком много файлов с именем {path.name}")


def download_attachment(url: str, cookie: str = "", max_bytes: int = 2_000_000) -> bytes:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept-Encoding": "identity",
    }
    if cookie:
        headers["Cookie"] = cookie
    ssl_ctx = ssl.create_default_context()
    ssl_ctx.check_hostname = False
    ssl_ctx.verify_mode = ssl.CERT_NONE
    request = Request(url, headers=headers)
    opener = urllib.request.build_opener(urllib.request.HTTPSHandler(context=ssl_ctx))
    with opener.open(request, timeout=30) as response:
        chunks: list[bytes] = []
        total = 0
        while True:
            chunk = response.read(min(65536, max_bytes - total + 1))
            if not chunk:
                break
            chunks.append(chunk)
            total += len(chunk)
            if total > max_bytes:
                raise ValueError(f"файл больше лимита {max_bytes} байт")
        return b"".join(chunks)


def download_attachment_via_browser(
    url: str,
    cdp_url: str,
    browser_profile_dir: str = "",
    max_bytes: int = 2_000_000,
) -> bytes:
    """Download a private attachment through the logged-in Chrome session."""
    from app.kwork_source import _ensure_chrome_cdp, _find_or_create_page, _send_cdp

    _ensure_chrome_cdp(cdp_url, "https://kwork.ru/projects", browser_profile_dir)
    page = _find_or_create_page(cdp_url, "https://kwork.ru/projects")

    import websocket

    ws = websocket.create_connection(page["webSocketDebuggerUrl"], timeout=45)
    try:
        _send_cdp(ws, "Network.enable", {})
        response = _send_cdp(ws, "Network.getCookies", {"urls": ["https://kwork.ru", url]})
    finally:
        ws.close()
    cookie_header = _cookie_header_from_cdp_cookies(response.get("result", {}).get("cookies", []), url)
    if not cookie_header:
        raise PermissionError("Chrome не отдал cookies Kwork. Проверь, что в отдельном Kwork Chrome выполнен вход.")
    return download_attachment(url, cookie=cookie_header, max_bytes=max_bytes)


def _download_with_fallback(
    ref: AttachmentRef,
    cookie: str,
    max_bytes: int,
    use_browser: bool,
    cdp_url: str,
    browser_profile_dir: str,
) -> bytes:
    try:
        content = download_attachment(ref.url, cookie=cookie, max_bytes=max_bytes)
        _validate_downloaded_content(ref, content)
        return content
    except Exception as direct_exc:
        if not use_browser:
            raise
        logger.info("Direct attachment download failed, trying Chrome session for %s: %s", ref.url, direct_exc)
        try:
            content = download_attachment_via_browser(
                ref.url,
                cdp_url=cdp_url,
                browser_profile_dir=browser_profile_dir,
                max_bytes=max_bytes,
            )
            _validate_downloaded_content(ref, content)
            return content
        except Exception as browser_exc:
            raise RuntimeError(
                f"прямое скачивание не прошло ({direct_exc}); Chrome-сессия тоже не скачала файл ({browser_exc}). "
                "Проверь, что в отдельном Kwork Chrome выполнен вход в аккаунт."
            ) from browser_exc


def extract_attachment_text(ref: AttachmentRef, content: bytes) -> str:
    return inspect_attachment(ref, content)[1]


def inspect_attachment(
    ref: AttachmentRef,
    content: bytes,
    max_bytes: int = 2_000_000,
    lead_context: str = "",
    deepseek_api_key: str = "",
    deepseek_model: str = "deepseek-chat",
    openrouter_api_key: str = "",
    openrouter_base_url: str = "https://openrouter.ai/api/v1",
    openrouter_vision_model: str = "",
) -> tuple[str, str]:
    ext = _extension(ref.url) or _extension(ref.label)
    if ext in TEXT_EXTENSIONS:
        return "скачан, прочитан", _decode_text(content)
    if ext == ".docx":
        docx_text = _extract_docx(content)
        if _docx_text_is_missing(docx_text):
            vision_text = describe_docx_with_openrouter(
                content,
                api_key=openrouter_api_key,
                model=openrouter_vision_model,
                base_url=openrouter_base_url,
            )
            if vision_text:
                return "скачан, vision прочитан", "DOCX без обычного текста. Vision:\n" + vision_text
        return "скачан, прочитан", docx_text
    if ext == ".pdf":
        return _extract_pdf(
            content,
            openrouter_api_key=openrouter_api_key,
            openrouter_base_url=openrouter_base_url,
            openrouter_vision_model=openrouter_vision_model,
        )
    if ext in IMAGE_EXTENSIONS:
        return _extract_image_ocr(
            content,
            ext,
            openrouter_api_key=openrouter_api_key,
            openrouter_base_url=openrouter_base_url,
            openrouter_vision_model=openrouter_vision_model,
        )
    if ext in ARCHIVE_EXTENSIONS:
        return _extract_archive(
            ref,
            content,
            max_bytes=max_bytes,
            lead_context=lead_context,
            deepseek_api_key=deepseek_api_key,
            deepseek_model=deepseek_model,
            openrouter_api_key=openrouter_api_key,
            openrouter_base_url=openrouter_base_url,
            openrouter_vision_model=openrouter_vision_model,
        )
    return "скачан, тип не поддержан", "Файл найден, но тип не поддержан для автоматического чтения."


def _attachment_kind(ref: AttachmentRef) -> str:
    ext = _extension(ref.url) or _extension(ref.label)
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in ARCHIVE_EXTENSIONS:
        return "archive"
    if ext == ".pdf":
        return "pdf"
    if ext in {".doc", ".docx"}:
        return "document"
    if ext in TEXT_EXTENSIONS:
        return "text"
    return ext.removeprefix(".") or "file"


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "DOCX найден, но python-docx недоступен."
    document = Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts) or "DOCX прочитан, но текст не найден."


def _docx_text_is_missing(text: str) -> bool:
    return not text.strip() or "текст не найден" in text.lower()


def _validate_downloaded_content(ref: AttachmentRef, content: bytes) -> None:
    ext = _extension(ref.url) or _extension(ref.label)
    if not content:
        raise ValueError("скачанный файл пустой")
    if ext in TEXT_EXTENSIONS:
        return
    if _looks_like_html(content):
        raise ValueError("Kwork вернул HTML-страницу вместо файла; нужен вход в аккаунт через Kwork Chrome")
    if ext == ".docx" and not zipfile.is_zipfile(io.BytesIO(content)):
        raise ValueError("скачанный файл не похож на DOCX")
    if ext == ".pdf" and not content.lstrip().startswith(b"%PDF"):
        raise ValueError("скачанный файл не похож на PDF")
    if ext == ".zip" and not zipfile.is_zipfile(io.BytesIO(content)):
        raise ValueError("скачанный файл не похож на ZIP-архив")


def _looks_like_html(content: bytes) -> bool:
    head = content[:512].lstrip().lower()
    return head.startswith(b"<!doctype html") or head.startswith(b"<html") or b"<body" in head[:256]


def _cookie_header_from_cdp_cookies(cookies: list[dict], url: str) -> str:
    host = (urlparse(url).hostname or "").lower()
    pairs: list[str] = []
    seen: set[str] = set()
    for cookie in cookies:
        name = str(cookie.get("name", "")).strip()
        value = str(cookie.get("value", ""))
        domain = str(cookie.get("domain", "")).lstrip(".").lower()
        if not name or name in seen:
            continue
        if domain and host != domain and not host.endswith(f".{domain}"):
            continue
        pairs.append(f"{name}={value}")
        seen.add(name)
    return "; ".join(pairs)


def _extract_pdf(
    content: bytes,
    openrouter_api_key: str = "",
    openrouter_base_url: str = "https://openrouter.ai/api/v1",
    openrouter_vision_model: str = "",
) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            return "скачан, текст не извлечен", "PDF найден, но библиотека для чтения PDF не установлена."
    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages[:5]:
        pages.append(page.extract_text() or "")
    extracted = "\n".join(text for text in pages if text.strip())
    if extracted.strip():
        return "скачан, прочитан", extracted
    try:
        ocr_text = _extract_pdf_ocr(content)
    except Exception as exc:
        vision_text = describe_pdf_with_openrouter(
            content,
            api_key=openrouter_api_key,
            model=openrouter_vision_model,
            base_url=openrouter_base_url,
        )
        if vision_text:
            return "скачан, vision прочитан", "PDF без текстового слоя. Vision:\n" + vision_text
        return "скачан, текст не извлечен", f"PDF без текстового слоя. OCR PDF не выполнен: {exc}"
    if ocr_text.strip():
        return "скачан, OCR прочитан", "PDF без текстового слоя. OCR:\n" + ocr_text
    vision_text = describe_pdf_with_openrouter(
        content,
        api_key=openrouter_api_key,
        model=openrouter_vision_model,
        base_url=openrouter_base_url,
    )
    if vision_text:
        return "скачан, vision прочитан", "PDF без текстового слоя. Vision:\n" + vision_text
    return "скачан, OCR не выполнен", "PDF без текстового слоя. OCR выполнился, но текст не найден."


def _extract_pdf_ocr(content: bytes, max_pages: int = 3) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF не установлен для OCR PDF") from exc
    document = fitz.open(stream=content, filetype="pdf")
    texts: list[str] = []
    try:
        for page_index in range(min(max_pages, document.page_count)):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_bytes = pixmap.tobytes("png")
            text = _run_tesseract_ocr(image_bytes, ".png")
            if text.strip():
                texts.append(f"Страница {page_index + 1}: {text.strip()}")
    finally:
        document.close()
    return "\n".join(texts)


def _extract_image_ocr(
    content: bytes,
    ext: str,
    openrouter_api_key: str = "",
    openrouter_base_url: str = "https://openrouter.ai/api/v1",
    openrouter_vision_model: str = "",
) -> tuple[str, str]:
    try:
        text = _run_tesseract_ocr(content, ext)
    except Exception as exc:
        vision_text = describe_image_with_openrouter(
            content,
            ext,
            api_key=openrouter_api_key,
            model=openrouter_vision_model,
            base_url=openrouter_base_url,
        )
        if vision_text:
            return "скачан, vision прочитан", vision_text
        return "скачан, OCR не выполнен", f"OCR не выполнен: {exc}"
    clean = _shorten(text, 2500)
    if clean:
        return "скачан, OCR прочитан", clean
    vision_text = describe_image_with_openrouter(
        content,
        ext,
        api_key=openrouter_api_key,
        model=openrouter_vision_model,
        base_url=openrouter_base_url,
    )
    if vision_text:
        return "скачан, vision прочитан", vision_text
    return "скачан, OCR не выполнен", "OCR выполнился, но текст на изображении не найден."


def _extract_archive(
    ref: AttachmentRef,
    content: bytes,
    max_bytes: int,
    lead_context: str,
    deepseek_api_key: str,
    deepseek_model: str,
    openrouter_api_key: str,
    openrouter_base_url: str,
    openrouter_vision_model: str,
) -> tuple[str, str]:
    ext = _extension(ref.url) or _extension(ref.label)
    if ext != ".zip":
        return "скачан, архив не открыт", "Автоматически открываются только ZIP-архивы; RAR/7Z пока нужно смотреть вручную."
    try:
        lines, selection = _read_zip_archive(
            ref,
            content,
            max_bytes=max_bytes,
            lead_context=lead_context,
            deepseek_api_key=deepseek_api_key,
            deepseek_model=deepseek_model,
            openrouter_api_key=openrouter_api_key,
            openrouter_base_url=openrouter_base_url,
            openrouter_vision_model=openrouter_vision_model,
        )
    except zipfile.BadZipFile:
        return "скачан, архив не открыт", "ZIP-архив поврежден или это не ZIP-файл."
    status = "скачан, архив открыт"
    if selection.used_ai:
        status += ", AI выбрала файлы" if selection.names else ", AI не выбрала файлы"
    return status, "\n".join(lines) or "ZIP-архив открыт, но подходящие файлы внутри не найдены."


def _read_zip_archive(
    ref: AttachmentRef,
    content: bytes,
    max_bytes: int,
    lead_context: str,
    deepseek_api_key: str,
    deepseek_model: str,
    openrouter_api_key: str,
    openrouter_base_url: str,
    openrouter_vision_model: str,
    max_entries: int = 8,
) -> tuple[list[str], ArchiveSelection]:
    lines: list[str] = []
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        infos = [info for info in archive.infolist() if not info.is_dir()]
        entries = tuple(_archive_entry_info(info) for info in infos)
        selection = select_archive_entries_with_deepseek(
            ref,
            entries,
            lead_context=lead_context,
            api_key=deepseek_api_key,
            model=deepseek_model,
            max_entries=max_entries,
        )
        selected_names = set(selection.names)
        if entries:
            lines.append("Состав архива: " + ", ".join(_format_archive_entry(entry) for entry in entries[:20]))
            if len(entries) > 20:
                lines.append(f"Еще файлов в архиве: {len(entries) - 20}, список сокращен.")
        if selection.used_ai:
            picked = ", ".join(selection.names) if selection.names else "ничего"
            reason = f"; причина: {selection.reason}" if selection.reason else ""
            lines.append(f"AI выбрала файлы: {picked}{reason}")
        else:
            picked = ", ".join(selection.names) if selection.names else "ничего"
            reason = f"; причина: {selection.reason}" if selection.reason else ""
            lines.append(f"Выбраны файлы для чтения: {picked}{reason}")

        info_by_name = {info.filename: info for info in infos}
        for name in selection.names:
            info = info_by_name.get(name)
            if info is None:
                continue
            name = info.filename
            if info.file_size > max_bytes:
                lines.append(f"{name}: пропущен, файл больше лимита {max_bytes} байт")
                continue
            with archive.open(info) as file:
                data = file.read(max_bytes + 1)
            if len(data) > max_bytes:
                lines.append(f"{name}: пропущен, файл больше лимита {max_bytes} байт")
                continue
            nested_ref = AttachmentRef(label=name, url=name)
            status, extracted = inspect_attachment(
                nested_ref,
                data,
                max_bytes=max_bytes,
                openrouter_api_key=openrouter_api_key,
                openrouter_base_url=openrouter_base_url,
                openrouter_vision_model=openrouter_vision_model,
            )
            nested_status = status.removeprefix("скачан, ")
            lines.append(f"{name}: {nested_status}\n{_shorten(extracted, 1500)}")
        skipped = [entry.name for entry in entries if entry.name not in selected_names]
        if skipped:
            lines.append(f"Не читались по выбору AI/fallback: {', '.join(skipped[:12])}")
            if len(skipped) > 12:
                lines.append(f"Еще пропущено файлов: {len(skipped) - 12}.")
    return lines, selection


def _archive_entry_info(info: zipfile.ZipInfo) -> ArchiveEntryInfo:
    return ArchiveEntryInfo(name=info.filename, size=int(info.file_size), kind=_archive_entry_kind(info.filename))


def _archive_entry_kind(name: str) -> str:
    ext = _extension(name)
    if ext in TEXT_EXTENSIONS:
        return "text"
    if ext == ".docx":
        return "document"
    if ext == ".pdf":
        return "pdf"
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in ARCHIVE_EXTENSIONS:
        return "archive"
    return ext.removeprefix(".") or "file"


def _format_archive_entry(entry: ArchiveEntryInfo) -> str:
    return f"{entry.name} ({entry.kind}, {entry.size} B)"


def select_archive_entries_with_deepseek(
    ref: AttachmentRef,
    entries: tuple[ArchiveEntryInfo, ...],
    lead_context: str = "",
    api_key: str = "",
    model: str = "deepseek-chat",
    max_entries: int = 8,
) -> ArchiveSelection:
    if not entries:
        return ArchiveSelection(names=(), used_ai=False, reason="архив пустой")
    fallback = _select_archive_entries_by_rules(entries, max_entries=max_entries)
    if not api_key:
        return fallback
    try:
        names, reason = _ask_deepseek_for_archive_entries(ref, entries, lead_context, api_key, model, max_entries)
    except Exception as exc:
        logger.warning("DeepSeek archive entry selection failed for %s: %s", ref.label, exc)
        return ArchiveSelection(names=fallback.names, used_ai=False, reason=f"AI не сработала, fallback: {fallback.reason}")
    valid_names = _valid_archive_names(names, entries, max_entries=max_entries)
    if not valid_names:
        return ArchiveSelection(names=fallback.names, used_ai=False, reason=f"AI не выбрала валидные файлы, fallback: {fallback.reason}")
    return ArchiveSelection(names=valid_names, used_ai=True, reason=_shorten(reason, 250))


def _ask_deepseek_for_archive_entries(
    ref: AttachmentRef,
    entries: tuple[ArchiveEntryInfo, ...],
    lead_context: str,
    api_key: str,
    model: str,
    max_entries: int,
) -> tuple[tuple[str, ...], str]:
    from openai import OpenAI

    entry_lines = "\n".join(f"- {entry.name} | {entry.kind} | {entry.size} B" for entry in entries[:60])
    prompt = (
        "Выбери файлы из ZIP-архива Kwork-заказа, которые нужно прочитать перед оценкой заказа и откликом.\n"
        "Выбирай только ТЗ, brief, описания, макеты, скриншоты, PDF/DOCX/TXT/HTML/CSV/JSON/XML. "
        "Не выбирай мусор, системные файлы, дубли и явно нерелевантные заметки. "
        f"Верни JSON строго вида {{\"read\": [\"filename\"], \"reason\": \"кратко\"}}. "
        f"Максимум файлов: {max_entries}.\n\n"
        f"Архив: {ref.label}\n"
        f"Контекст заказа:\n{_shorten(lead_context, 2500) or 'нет контекста'}\n\n"
        f"Файлы:\n{entry_lines}"
    )
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com/v1")
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "Ты аккуратный помощник, выбираешь только полезные файлы ТЗ внутри архива."},
            {"role": "user", "content": prompt},
        ],
        temperature=0,
        max_tokens=500,
    )
    content = response.choices[0].message.content or ""
    payload = _parse_json_object(content)
    raw_names = payload.get("read", [])
    if not isinstance(raw_names, list):
        raw_names = []
    names = tuple(str(name).strip() for name in raw_names if str(name).strip())
    reason = str(payload.get("reason", "")).strip()
    return names, reason


def _parse_json_object(raw: str) -> dict:
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end < start:
        raise ValueError("DeepSeek response does not contain JSON object")
    payload = json.loads(raw[start : end + 1])
    if not isinstance(payload, dict):
        raise ValueError("DeepSeek response JSON must be an object")
    return payload


def _valid_archive_names(names: tuple[str, ...], entries: tuple[ArchiveEntryInfo, ...], max_entries: int) -> tuple[str, ...]:
    allowed = {entry.name for entry in entries}
    result: list[str] = []
    for name in names:
        if name in allowed and name not in result:
            result.append(name)
        if len(result) >= max_entries:
            break
    return tuple(result)


def _select_archive_entries_by_rules(entries: tuple[ArchiveEntryInfo, ...], max_entries: int) -> ArchiveSelection:
    readable = [entry for entry in entries if entry.kind in {"text", "document", "pdf", "image"}]
    if not readable:
        return ArchiveSelection(names=tuple(entry.name for entry in entries[:max_entries]), used_ai=False, reason="поддержанных файлов не найдено")
    scored = sorted(readable, key=_archive_entry_priority)
    names = tuple(entry.name for entry in scored[:max_entries])
    return ArchiveSelection(names=names, used_ai=False, reason="выбраны похожие на ТЗ файлы")


def _archive_entry_priority(entry: ArchiveEntryInfo) -> tuple[int, int, str]:
    lowered = entry.name.lower()
    if any(part in lowered for part in IMPORTANT_ARCHIVE_NAME_PARTS):
        name_score = 0
    else:
        name_score = 1
    kind_score = {"document": 0, "pdf": 1, "text": 2, "image": 3}.get(entry.kind, 9)
    return (name_score, kind_score, entry.name)


def _run_tesseract_ocr(content: bytes, ext: str) -> str:
    command = _tesseract_command()
    if not command.exists():
        raise FileNotFoundError(f"Tesseract не найден: {command}")
    suffix = ext if ext.startswith(".") else f".{ext}"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as image_file:
        image_file.write(content)
        image_path = Path(image_file.name)
    try:
        result = subprocess.run(
            [str(command), str(image_path), "stdout", "-l", "rus+eng", "--psm", "6"],
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=45,
        )
    finally:
        image_path.unlink(missing_ok=True)
    if result.returncode != 0:
        error = (result.stderr or result.stdout or "неизвестная ошибка").strip()
        raise RuntimeError(error)
    return result.stdout.strip()


def describe_image_with_openrouter(
    content: bytes,
    extension: str,
    api_key: str,
    model: str,
    base_url: str,
    timeout_seconds: float = 45.0,
) -> str:
    """Describe a technical screenshot only when optional OpenRouter vision is configured."""
    if not api_key.strip() or not model.strip():
        return ""
    mime_type = _image_mime_type(extension)
    image_data = b64encode(content).decode("ascii")
    try:
        from openai import OpenAI

        client = OpenAI(
            api_key=api_key,
            base_url=base_url.strip() or "https://openrouter.ai/api/v1",
            timeout=timeout_seconds,
        )
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты читаешь изображение, приложенное к Kwork-заказу. "
                        "Опиши только наблюдаемые элементы ТЗ, интерфейса и требования. "
                        "Не выдумывай детали, не пиши цену и не обращайся к заказчику."
                    ),
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Кратко опиши, что важно реализовать или учесть на этом изображении.",
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{mime_type};base64,{image_data}"},
                        },
                    ],
                },
            ],
            temperature=0.1,
            max_tokens=700,
        )
    except Exception as exc:
        logger.warning("OpenRouter vision failed for %s: %s", extension, exc)
        return ""
    return _shorten((response.choices[0].message.content or "").strip(), 2500)


def describe_pdf_with_openrouter(
    content: bytes,
    api_key: str,
    model: str,
    base_url: str,
    timeout_seconds: float = 45.0,
) -> str:
    """Render the first scan pages and ask OpenRouter vision only after local OCR fails."""
    if not api_key.strip() or not model.strip():
        return ""
    try:
        import fitz

        document = fitz.open(stream=content, filetype="pdf")
        descriptions: list[str] = []
        try:
            for page_index in range(min(2, document.page_count)):
                pixmap = document.load_page(page_index).get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                description = describe_image_with_openrouter(
                    pixmap.tobytes("png"),
                    ".png",
                    api_key=api_key,
                    model=model,
                    base_url=base_url,
                    timeout_seconds=timeout_seconds,
                )
                if description:
                    descriptions.append(f"Страница {page_index + 1}: {description}")
        finally:
            document.close()
        return _shorten("\n".join(descriptions), 2500)
    except Exception as exc:
        logger.warning("OpenRouter vision could not render PDF: %s", exc)
        return ""


def describe_docx_with_openrouter(
    content: bytes,
    api_key: str,
    model: str,
    base_url: str,
    timeout_seconds: float = 45.0,
) -> str:
    """Describe embedded DOCX screenshots when ordinary text extraction finds nothing."""
    if not api_key.strip() or not model.strip():
        return ""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            image_names = [
                name
                for name in archive.namelist()
                if name.lower().startswith("word/media/") and _extension(name) in IMAGE_EXTENSIONS
            ]
            descriptions: list[str] = []
            for name in image_names[:2]:
                description = describe_image_with_openrouter(
                    archive.read(name),
                    _extension(name),
                    api_key=api_key,
                    model=model,
                    base_url=base_url,
                    timeout_seconds=timeout_seconds,
                )
                if description:
                    descriptions.append(f"{Path(name).name}: {description}")
        return _shorten("\n".join(descriptions), 2500)
    except Exception as exc:
        logger.warning("OpenRouter vision could not inspect DOCX images: %s", exc)
        return ""


def _image_mime_type(extension: str) -> str:
    normalized = extension.lower().lstrip(".")
    if normalized in {"jpg", "jpeg"}:
        return "image/jpeg"
    if normalized == "webp":
        return "image/webp"
    if normalized == "gif":
        return "image/gif"
    return "image/png"


def _tesseract_command() -> Path:
    configured = os.getenv("TESSERACT_CMD", "").strip()
    if configured:
        path = Path(configured)
        if path.drive.upper() != "D:":
            raise RuntimeError("TESSERACT_CMD должен указывать на D: диск")
        return path
    return DEFAULT_TESSERACT_CMD


def _extension(value: str) -> str:
    path = urlparse(value).path if re.match(r"https?://", value, re.IGNORECASE) else value
    match = re.search(r"(\.[A-Za-z0-9]+)$", path)
    return match.group(1).lower() if match else ""


def _shorten(text: str, limit: int) -> str:
    text = " ".join(text.split())
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"
